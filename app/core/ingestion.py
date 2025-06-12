from pathlib import Path
from typing import List, Optional
import pytesseract
from PIL import Image
import docx
import fitz  # PyMuPDF
import requests
from bs4 import BeautifulSoup
from langchain.schema import Document
from app.core.engine_singleton import get_rag_engine
from app.core.logger import setup_logger
import aiohttp
import asyncio
import os
from langchain.text_splitter import RecursiveCharacterTextSplitter
import datetime

logger = setup_logger(__name__)

def get_text_splitter() -> RecursiveCharacterTextSplitter:
    """Get a configured text splitter instance."""
    return RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        length_function=len,
        separators=["\n\n", "\n", " ", ""]
    )

async def split_text_into_documents(text: str, metadata: dict) -> List[Document]:
    """Split text into chunks and create documents."""
    text_splitter = get_text_splitter()
    texts = text_splitter.split_text(text)
    
    documents = []
    for i, chunk in enumerate(texts):
        doc = Document(
            page_content=chunk,
            metadata={
                **metadata,
                "chunk": i,
                "chunk_size": len(chunk)
            }
        )
        documents.append(doc)
    
    return documents

async def process_text_file(file_path: Path) -> List[Document]:
    """Process a text file and return a list of documents."""
    logger.info(f"Processing text file: {file_path}")
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
            logger.debug(f"Read {len(text)} characters from {file_path}")
        
        metadata = {
            "file": file_path.name,
            "file_type": "txt",
            "source": str(file_path)
        }
        
        documents = await split_text_into_documents(text, metadata)
        logger.info(f"Successfully processed text file: {file_path} into {len(documents)} chunks")
        return documents
    except Exception as e:
        logger.error(f"Failed to process text file {file_path}: {str(e)}", exc_info=True)
        raise

async def process_pdf_file(file_path: Path) -> List[Document]:
    """Process a PDF file and return a list of documents."""
    logger.info(f"Processing PDF file: {file_path}")
    try:
        documents = []
        pdf = fitz.open(file_path)
        
        for page_num in range(len(pdf)):
            page = pdf[page_num]
            text = page.get_text()
            
            if text.strip():
                metadata = {
                    "file": file_path.name,
                    "file_type": "pdf",
                    "source": str(file_path),
                    "page": page_num + 1
                }
                
                page_docs = await split_text_into_documents(text, metadata)
                documents.extend(page_docs)
                logger.debug(f"Processed page {page_num + 1} into {len(page_docs)} chunks")
        
        logger.info(f"Successfully processed PDF file: {file_path} into {len(documents)} chunks across {len(pdf)} pages")
        return documents
    except Exception as e:
        logger.error(f"Failed to process PDF file {file_path}: {str(e)}", exc_info=True)
        raise

async def process_docx_file(file_path: Path) -> List[Document]:
    """Process a DOCX file and return a list of documents."""
    logger.info(f"Processing DOCX file: {file_path}")
    try:
        doc = docx.Document(file_path)
        text = "\n\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])
        logger.debug(f"Extracted {len(text)} characters from DOCX file")
        
        metadata = {
            "file": file_path.name,
            "file_type": "docx",
            "source": str(file_path)
        }
        
        documents = await split_text_into_documents(text, metadata)
        logger.info(f"Successfully processed DOCX file: {file_path} into {len(documents)} chunks")
        return documents
    except Exception as e:
        logger.error(f"Failed to process DOCX file {file_path}: {str(e)}", exc_info=True)
        raise

async def process_image_file(file_path: Path) -> List[Document]:
    """Process an image file using OCR and return a list of documents."""
    logger.info(f"Processing image file: {file_path}")
    try:
        image = Image.open(file_path)
        text = pytesseract.image_to_string(image)
        
        if text.strip():
            document = Document(
                page_content=text,
                metadata={
                    "file": file_path.name,
                    "file_type": "image",
                    "source": str(file_path)
                }
            )
            logger.info(f"Successfully processed image file: {file_path}")
            return [document]
        else:
            logger.warning(f"No text extracted from image: {file_path}")
            return []
    except Exception as e:
        logger.error(f"Failed to process image file {file_path}: {str(e)}", exc_info=True)
        raise

async def process_web_page(url: str) -> List[Document]:
    """Process a web page and return a list of documents."""
    logger.info(f"Processing web page: {url}")
    try:
        response = requests.get(url)
        response.raise_for_status()
        
        soup = BeautifulSoup(response.text, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
        
        text = soup.get_text()
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = ' '.join(chunk for chunk in chunks if chunk)
        
        document = Document(
            page_content=text,
            metadata={"source": url, "file_type": "web"}
        )
        logger.info(f"Successfully processed web page: {url}")
        return [document]
    except Exception as e:
        logger.error(f"Failed to process web page {url}: {str(e)}", exc_info=True)
        raise

async def process_file(file_path: Path) -> List[Document]:
    """Process a file based on its extension."""
    logger.info(f"Processing file: {file_path}")
    try:
        file_processors = {
            ".txt": process_text_file,
            ".md": process_text_file,
            ".pdf": process_pdf_file,
            ".docx": process_docx_file,
            ".png": process_image_file,
            ".jpg": process_image_file,
            ".jpeg": process_image_file
        }
        
        extension = file_path.suffix.lower()
        processor = file_processors.get(extension)
        
        if processor:
            documents = await processor(file_path)
            logger.info(f"Successfully processed file: {file_path}")
            return documents
        else:
            logger.warning(f"Unsupported file type: {extension}")
            raise ValueError(f"Unsupported file type: {extension}")
    except Exception as e:
        logger.error(f"Failed to process file {file_path}: {str(e)}", exc_info=True)
        raise

async def ingest_file(file_path: Path) -> bool:
    """Ingest a file into the vector store."""
    logger.info(f"Starting ingestion for file: {file_path}")
    try:
        # Process the file
        documents = await process_file(file_path)
        
        if documents:
            logger.info(f"Successfully extracted {len(documents)} chunks from {file_path}")
            
            # Get RAG engine instance
            rag_engine = get_rag_engine()
            
            # Add file metadata
            metadata = {
                "file": file_path.name,
                "source": str(file_path),
                "ingestion_time": str(datetime.datetime.now()),
                "file_type": file_path.suffix.lower(),
                "file_size": os.path.getsize(file_path),
                "total_chunks": len(documents)
            }
            
            # Add documents to the vector store
            await rag_engine.add_documents(
                documents=documents,
                source_type="file",
                metadata=metadata
            )
            
            logger.info(f"Successfully ingested file: {file_path} with {len(documents)} chunks")
            return True
        else:
            logger.warning(f"No documents extracted from file: {file_path}")
            return False
    except Exception as e:
        logger.error(f"Error ingesting file {file_path}: {str(e)}", exc_info=True)
        return False

async def ingest_web_page(url: str) -> bool:
    """Ingest content from a web page."""
    logger.info(f"Starting ingestion for web page: {url}")
    try:
        # Get RAG engine instance
        rag_engine = get_rag_engine()
        
        # Clean up any previous web results
        rag_engine._cleanup_web_results()
        
        # Fetch the web page content
        async with aiohttp.ClientSession() as session:
            async with session.get(url) as response:
                if response.status != 200:
                    logger.error(f"Failed to fetch URL {url}, status code: {response.status}")
                    return False
                
                html_content = await response.text()

        # Parse HTML content
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()

        # Get text content
        text = soup.get_text()
        
        # Clean up text
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = ' '.join(chunk for chunk in chunks if chunk)

        # Split text into chunks
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
        )
        
        chunks = text_splitter.split_text(text)
        
        # Create documents
        documents = []
        doc_ids = []
        for i, chunk in enumerate(chunks):
            doc_id = f"web_{url}_{i}"
            doc_ids.append(doc_id)
            doc = Document(
                page_content=chunk,
                metadata={
                    "source": "web",
                    "url": url,
                    "chunk": i,
                    "title": soup.title.string if soup.title else url,
                    "id": doc_id
                }
            )
            documents.append(doc)

        # Add documents with IDs
        rag_engine.vector_store.add_documents(documents, ids=doc_ids)
        # Update tracking set
        rag_engine.current_web_ids.update(doc_ids)
        
        logger.info(f"Successfully ingested {len(documents)} chunks from {url}")
        return True

    except Exception as e:
        logger.error(f"Error ingesting web page {url}: {str(e)}", exc_info=True)
        return False 