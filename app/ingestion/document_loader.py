from typing import List, Dict, Any
from fastapi import UploadFile
from pathlib import Path
import fitz  # PyMuPDF
from docx import Document as DocxDocument
import markdown
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
import os
import pytesseract
from PIL import Image
import tempfile

async def process_document(
    file: UploadFile,
    metadata: Dict[str, Any]
) -> List[Document]:
    """
    Process a document file and convert it to a list of Document objects.
    
    Args:
        file: The uploaded document file
        metadata: Additional metadata to attach to the documents
        
    Returns:
        List of Document objects
    """
    # Create text splitter
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        length_function=len,
        separators=["\n\n", "\n", " ", ""]
    )
    
    # Get file extension
    file_extension = Path(file.filename).suffix.lower()
    
    # Save file to temporary location
    with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as temp_file:
        content = await file.read()
        temp_file.write(content)
        temp_file_path = temp_file.name
    
    try:
        # Process based on file type
        if file_extension in ['.pdf']:
            documents = await process_pdf(temp_file_path, text_splitter, metadata)
        elif file_extension in ['.docx', '.doc']:
            documents = await process_docx(temp_file_path, text_splitter, metadata)
        elif file_extension in ['.txt']:
            documents = await process_text(temp_file_path, text_splitter, metadata)
        elif file_extension in ['.md']:
            documents = await process_markdown(temp_file_path, text_splitter, metadata)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        return documents
    finally:
        # Clean up temporary file
        os.unlink(temp_file_path)

async def process_pdf(
    file_path: str,
    text_splitter: RecursiveCharacterTextSplitter,
    metadata: Dict[str, Any]
) -> List[Document]:
    """Process a PDF file."""
    documents = []
    
    # Open PDF
    pdf_document = fitz.open(file_path)
    
    for page_num in range(len(pdf_document)):
        page = pdf_document[page_num]
        
        # Try to extract text
        text = page.get_text()
        
        # If no text found, try OCR
        if not text.strip():
            # Convert page to image
            pix = page.get_pixmap()
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            
            # Perform OCR
            text = pytesseract.image_to_string(img)
        
        if text.strip():
            # Create page-specific metadata
            page_metadata = {
                **metadata,
                "page_number": page_num + 1,
                "total_pages": len(pdf_document)
            }
            
            # Split text into chunks
            chunks = text_splitter.split_text(text)
            
            # Create documents
            for chunk in chunks:
                documents.append(
                    Document(
                        page_content=chunk,
                        metadata=page_metadata
                    )
                )
    
    return documents

async def process_docx(
    file_path: str,
    text_splitter: RecursiveCharacterTextSplitter,
    metadata: Dict[str, Any]
) -> List[Document]:
    """Process a DOCX file."""
    documents = []
    
    # Load document
    doc = DocxDocument(file_path)
    
    # Extract text from paragraphs
    full_text = "\n\n".join([paragraph.text for paragraph in doc.paragraphs])
    
    # Split text into chunks
    chunks = text_splitter.split_text(full_text)
    
    # Create documents
    for chunk in chunks:
        documents.append(
            Document(
                page_content=chunk,
                metadata=metadata
            )
        )
    
    return documents

async def process_text(
    file_path: str,
    text_splitter: RecursiveCharacterTextSplitter,
    metadata: Dict[str, Any]
) -> List[Document]:
    """Process a text file."""
    documents = []
    
    # Read file
    with open(file_path, 'r', encoding='utf-8') as file:
        text = file.read()
    
    # Split text into chunks
    chunks = text_splitter.split_text(text)
    
    # Create documents
    for chunk in chunks:
        documents.append(
            Document(
                page_content=chunk,
                metadata=metadata
            )
        )
    
    return documents

async def process_markdown(
    file_path: str,
    text_splitter: RecursiveCharacterTextSplitter,
    metadata: Dict[str, Any]
) -> List[Document]:
    """Process a Markdown file."""
    documents = []
    
    # Read file
    with open(file_path, 'r', encoding='utf-8') as file:
        text = file.read()
    
    # Convert markdown to plain text
    html = markdown.markdown(text)
    # TODO: Use BeautifulSoup to clean HTML and extract text if needed
    
    # Split text into chunks
    chunks = text_splitter.split_text(text)
    
    # Create documents
    for chunk in chunks:
        documents.append(
            Document(
                page_content=chunk,
                metadata=metadata
            )
        )
    
    return documents 